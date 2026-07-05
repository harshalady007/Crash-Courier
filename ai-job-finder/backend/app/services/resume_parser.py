"""Resume text extraction (PDF/DOCX) and heuristic section parsing.

Extraction: PyMuPDF for PDF, python-docx for DOCX.
Parsing: section-header segmentation + targeted regexes. Deliberately heuristic —
production can layer LLM structured extraction on top (docs/08) with this as fallback.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime

from .skill_extractor import extract_skills, tools_subset

# ---------------------------------------------------------------- text extraction

class UnsupportedFileError(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> str:
    """Extract plain text from a PDF or DOCX upload."""
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return _extract_pdf(content)
    if lowered.endswith(".docx"):
        return _extract_docx(content)
    raise UnsupportedFileError("Only .pdf and .docx resumes are supported.")


def _extract_pdf(content: bytes) -> str:
    import fitz  # PyMuPDF

    with fitz.open(stream=content, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _extract_docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ---------------------------------------------------------------- section parsing

SECTION_HEADERS = {
    "education": r"education|academic|qualification",
    "skills": r"skills|technical skills|technologies|tech stack|core competencies",
    "projects": r"projects|personal projects|academic projects",
    "experience": r"experience|work experience|employment|internship|professional experience",
    "certifications": r"certifications?|courses|licenses",
    "objective": r"objective|summary|profile|career goal|about",
}

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{3,5}\)?[\s-]?)?\d{3}[\s-]?\d{4}")
DEGREE_RE = re.compile(
    r"(b\.?\s?tech|b\.?e\.?|b\.?sc|bca|m\.?\s?tech|m\.?sc|mca|mba|bachelor|master|ph\.?d|diploma)",
    re.IGNORECASE,
)
YEAR_RE = re.compile(r"(19|20)\d{2}")
ROLE_HINT_RE = re.compile(
    r"(data analyst|data scientist|machine learning|ai engineer|software engineer|"
    r"python developer|business analyst|automation engineer|web developer|full stack)",
    re.IGNORECASE,
)
DATE_RANGE_RE = re.compile(
    r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(\d{4})\s*(?:-|–|to)\s*"
    r"(?:(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(\d{4})|present|current|now)",
    re.IGNORECASE,
)
YEARS_PHRASE_RE = re.compile(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?experience", re.IGNORECASE)

MONTHS = {m: i + 1 for i, m in enumerate(
    ["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"]
)}


@dataclass
class ParsedResume:
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    skills: list[str] = field(default_factory=list)
    tools: list[str] = field(default_factory=list)
    education: list[dict] = field(default_factory=list)
    projects: list[dict] = field(default_factory=list)
    experience: list[dict] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    target_roles: list[str] = field(default_factory=list)
    experience_years: float = 0.0


def split_sections(text: str) -> dict[str, str]:
    """Segment resume text by recognized section headers. 'header' = short line matching a known pattern."""
    lines = text.splitlines()
    sections: dict[str, list[str]] = {"_top": []}
    current = "_top"
    for line in lines:
        stripped = line.strip().strip(":").lower()
        matched = None
        if stripped and len(stripped) <= 40:
            for key, pattern in SECTION_HEADERS.items():
                if re.fullmatch(rf"(?:{pattern})\s*", stripped, re.IGNORECASE):
                    matched = key
                    break
        if matched:
            current = matched
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items()}


def _parse_name(top: str, full_text: str) -> str | None:
    """First short, letters-only line that isn't contact info is usually the name."""
    for line in (top or full_text).splitlines()[:6]:
        candidate = line.strip()
        if not candidate or EMAIL_RE.search(candidate) or any(ch.isdigit() for ch in candidate):
            continue
        if 2 <= len(candidate.split()) <= 5 and len(candidate) <= 50:
            return candidate.title() if candidate.isupper() else candidate
    return None


def _parse_education(section: str) -> list[dict]:
    items = []
    for line in section.splitlines():
        degree_match = DEGREE_RE.search(line)
        if degree_match:
            year_match = YEAR_RE.search(line)
            items.append({
                "degree": degree_match.group(0).strip(),
                "institution": re.sub(r"\s{2,}", " ", line).strip()[:120],
                "year": year_match.group(0) if year_match else "",
            })
    return items[:5]


def _parse_list_items(section: str) -> list[dict]:
    """Bullet/paragraph items → [{title, description}].

    Heuristic: a short line that doesn't end with a period starts a new item;
    following longer lines are its description.
    """
    items: list[dict] = []
    current_title: str | None = None
    current_desc: list[str] = []

    def flush() -> None:
        if current_title:
            items.append({"title": current_title, "description": " ".join(current_desc)[:500]})

    for line in section.splitlines():
        stripped = line.strip("•-*● \t")
        if not stripped:
            continue
        looks_like_title = len(stripped) <= 70 and not stripped.endswith(".")
        if looks_like_title and (current_title is None or current_desc):
            flush()
            current_title, current_desc = stripped[:120], []
        else:
            current_desc.append(stripped)
    flush()
    return items[:10]


def _estimate_experience_years(text: str, experience_section: str) -> float:
    """Sum parsed date ranges; fall back to 'X years of experience' phrases."""
    total_months = 0.0
    now = datetime.now()
    for m in DATE_RANGE_RE.finditer(experience_section or text):
        start = datetime(int(m.group(2)), MONTHS[m.group(1).lower()[:3]], 1)
        if m.group(4):
            end = datetime(int(m.group(4)), MONTHS[m.group(3).lower()[:3]], 1)
        else:
            end = now
        months = max(0.0, (end - start).days / 30.44)
        total_months += min(months, 120)  # cap absurd ranges
    if total_months == 0:
        phrase = YEARS_PHRASE_RE.search(text)
        if phrase:
            return min(float(phrase.group(1)), 40.0)
    return round(total_months / 12, 1)


def parse_resume(text: str) -> ParsedResume:
    sections = split_sections(text)
    top = sections.get("_top", "")

    skills = extract_skills(text)
    email_match = EMAIL_RE.search(text)
    phone_match = PHONE_RE.search(top or text)

    certifications = [
        line.strip("•-*● \t")
        for line in sections.get("certifications", "").splitlines()
        if line.strip("•-*● \t")
    ][:10]

    target_roles = sorted({m.group(0).title() for m in ROLE_HINT_RE.finditer(sections.get("objective", "") or top)})

    experience_items = _parse_list_items(sections.get("experience", ""))

    return ParsedResume(
        name=_parse_name(top, text),
        email=email_match.group(0) if email_match else None,
        phone=phone_match.group(0).strip() if phone_match else None,
        skills=skills,
        tools=tools_subset(skills),
        education=_parse_education(sections.get("education", "")),
        projects=_parse_list_items(sections.get("projects", "")),
        experience=experience_items,
        certifications=certifications,
        target_roles=target_roles,
        experience_years=_estimate_experience_years(text, sections.get("experience", "")),
    )
